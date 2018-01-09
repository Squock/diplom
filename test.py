from django.shortcuts import render
from django.contrib import auth
from django.contrib.auth.decorators import login_required
from django.views.generic import ListView
from django.contrib.auth import authenticate
from django.conf import settings
from django.shortcuts import redirect
from django.core.context_processors import csrf
from django.contrib.auth.forms import UserCreationForm
from django.shortcuts import render_to_response
from docx import *
from docx import Document
from django.contrib.auth.models import User
from docx.shared import Inches
from blog.models import DocumentType, Filling, userdata, worker, UploadFileForm, UserDoc, orgdata
from django.http import HttpResponseRedirect, HttpResponse
import os
from io import StringIO
from django.views.decorators.csrf import requires_csrf_token
from django.core.files import File
import glob
import zipfile
from collections import defaultdict

Dictionary = {'name':'Имя',
'surname':'Фамилия',
'second_name':'Отчество',
'position':'Должность',
'inn':'inn',
'ogrn':'ogrn',
'orgname':'orgname',
'address':'address',
'orgcity': 'city'}
dict_commision = defaultdict(list)
dict_commision['name']
idd = []
def logout(request):
    auth.logout(request)
    # Перенаправление на страницу.
    return HttpResponseRedirect("/")

@login_required
def contact(request):
    return render(request, 'blog/contact.html', {})

@login_required
def cabinet(request):
    message = " "

    current_user = request.user.id
    count = False
    data = userdata.objects.all()
    for f in data:
        #print("f.user1",f.user_id)
        #print(f.second_name)
        if(f.user_id == current_user):
            cur_user = userdata.objects.get(pk=f.id)
            count = True
    if(count == False):
        print("phh")
        return HttpResponseRedirect("/userdata")
        """else:
            return HttpResponseRedirect("/userdata")"""
    #cur_userdata = [ud for ud in cur_user.userdata_set.all()]
    #cur_user = userdata.objects.get(pk=request.user.id)

    if request.method == "POST":
        cur_user.first_name = request.POST["first_name"]
        cur_user.last_name = request.POST["last_name"]
        cur_user.email = request.POST["email"]
        cur_user.number = request.POST["number"]
        cur_user.second_name = request.POST["second_name"]
        cur_user.save()
        message = "Данные успешно сохранены"
#     print(type(cur_user))
#     all_entries = userdata.objects.all()
#     current_user = request.user.id
# #        user_id = models.ForeignKey(User)
#     for a in userdata.objects.all():
#         a.user_id
#         a.id
#         print(a.user_id)
#         if a.user_id == current_user:
#             return render(request, 'blog/cabinet.html', {"all_entries":all_entries})
    return render(request, 'blog/cabinet.html', {"user":cur_user, "message":message})

@login_required
def pass_set(request):
    message = ""
    if request.method == "POST":
        cur_user = request.user
        u = User.objects.get(username=request.user)
        password = request.POST["password2"]
        password2 = request.POST["password"]
        if(password == password2):
            u.set_password(password)
            u.save()
            print(u)
        else:
            message = "Данные не совпадают"
    return render(request, 'blog/pass_set.html', {"message":message})

@login_required
def post_list(request):
    items1 = UserDoc.objects.all()
    current_user = request.user.id
    filename = UploadFileForm.objects.all()
    for f in filename:
        docfile = f.file
    for f in items1:
        if(f.user_id == current_user):
            #cur_user = userdata.objects.get(pk=f.id)
            items = str(f.userFile)
        else:
            items = None

    if request.method == 'POST':
        #return render(request, 'blog/post_list.html', {'items': items})

        return HttpResponse(docfile, mimetype='application/octet-stream')
    # args['username'] = auth.get_user(request)
#    if not request.user.is_authenticated():
#        return render(request, 'blog/post_list.html', args)
#    else:
#        return render(request, 'registration/login.html', {})
    return render(request, 'blog/post_list.html', {})
    #return render(request, 'blog/post_list.html', {"items":items})

def download(request):
    global doctype, zipp
    userdoc = UserDoc.objects.all()
    test = UploadFileForm.objects.all()
    for u in userdoc:
        if(request.user.id == u.user_id):
            print(u.user_id)
            for t in test:
                #docfiles = open('C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\' + str(request.user.id) + t.title + '.docx', 'rb')
                zipfile = open("C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\document" + str(request.user.id) + ".zip", 'rb')
                #response = HttpResponse(content=docfiles)
                response = HttpResponse(content=zipfile)
                #response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                response['Content-Type'] = 'application/zip'
                #response['Content-Disposition'] = "attachment; filename=test.docx"
                response['Content-Disposition'] = "attachment; filename=document.zip"
                #if doctype == t.title:
                #print(doctype)
                #return response

class DocumentTypeListView(ListView):
    #username = request.POST['username']
    #password = request.POST['password']
    model = DocumentType

def login(request):
    username = request.POST['username']
    password = request.POST['password']
    user = auth.authenticate(username=username, password=password)
    if user is not None and user.is_active:
        # Правильный пароль и пользователь "активен"
        auth.login(request, user)
        # Перенаправление на "правильную" страницу
        return HttpResponseRedirect("/account/loggedin/")
    else:
        # Отображение страницы с ошибкой
        return HttpResponseRedirect("/account/invalid/")


def registrationView(request): #Функция регистрации при помощи использование форм Django
    args ={}
    args.update(csrf(request))
    args['form'] = UserCreationForm()
    if request.method == 'POST':
        newuser_form = UserCreationForm(request.POST)
        if newuser_form.is_valid():
            newuser_form.save()
            newuser = auth.authenticate(username=newuser_form.cleaned_data['username'], password=newuser_form.cleaned_data['password2'])
        else:
            args['form'] = newuser_form
    return render_to_response('registration/registration.html', args)

@requires_csrf_token
def userView(request):
    if request.method == 'POST':
        inputFirstName = request.POST.get('first_name')
        inputLastName = request.POST.get('last_name')
        inputSecondName = request.POST.get('second_name')
        inputEmail = request.POST.get('email')
        number = request.POST.get('number')
        b = userdata.objects.create(user=request.user, first_name=inputFirstName, last_name=inputLastName, second_name=inputSecondName, email=inputEmail, number=number)
        b.save()
        c = {}
        #c.update(csrf(request))
        return HttpResponseRedirect("/")
    return render(request, 'registration/user.html',{})


@login_required
def fillingView(request):
    #cur_user = Filling.objects.get(pk=request.user.id)
    #print(cur_user)
    global idd
    find = False
    current_user = request.user.id
    count = False
    data = Filling.objects.all()
    orgdata_doc = orgdata.objects.all()
    message = ""
    for f in data:
        if(current_user == f.author_id_id):
            cur_user = Filling.objects.get(pk=f.id)
            #if(f.author_id_id == current_user):
            count = True
            find = True
            test = f.id
    for o in orgdata_doc:
        if(current_user == o.user_id):
            user_data = orgdata.objects.get(pk=o.id)
            org_id = o.id
    if(count == False):
        find = False
        cur_user = None
        #return HttpResponseRedirect("/worker_change")
        #return render_to_response('blog/filling_list_change.html')
    userdoc = UserDoc.objects.all()
    finded = False
    for u in userdoc:
        if(u.user_id == current_user):
            #cur_user = UserDoc.objects.get(pk=u.id)
            finded = True
            uid = u.id

    if request.method =='POST':
        inputInn = request.POST.get('inn')
        inputOgrn = request.POST.get('ogrn')
        inputCity = request.POST.get('city')
        inputAddress = request.POST.get('address')
        inputName = request.POST.get('name')
        inputChiefName = request.POST.get('chief_name')
        inputChiefSurname = request.POST.get('chief_surname')
        inputChiefSecondname = request.POST.get('chief_secondname')
        inputChiefFullposition = request.POST.get('chief_fullposition')
        message = "Данные успешно сохранены"
        #cur_user.save()
#        user_id = models.ForeignKey(User)
        """for a in Filling.objects.all():

            a.author_id_id
            a.id
            print("cur_id", current_user)
            print("cur_id", current_user, "a.id",a.id, "auth_id",a.author_id_id)"""

        if(find == True):
            idd = f.id
            b = Filling(id=test, author_id=request.user, inn=inputInn, ogrn=inputOgrn, name=inputName, city=inputCity, address=inputAddress, chief_name=inputChiefName, chief_surname=inputChiefSurname, chief_secondname=inputChiefSecondname, chief_fullposition=inputChiefFullposition)
            c = orgdata.objects.create(doc_id=f.id, user_id=request.user.id)
            #c = orgdata(id=org_id, doc_id=f.id, user_id=request.user.id)
            c.save()
            b.save()
        else:
            b = Filling.objects.create(author_id=request.user, inn=inputInn, ogrn=inputOgrn, name=inputName, city=inputCity, address=inputAddress, chief_name=inputChiefName, chief_surname=inputChiefSurname, chief_secondname=inputChiefSecondname, chief_fullposition=inputChiefFullposition)
            c = orgdata.objects.create(doc_id=f.id, user_id=request.user.id)
            c.save()
            #c = orgdata.objects.create(doc_id = request.)
            b.save()

        #
        #doc = open('C:\myproject\Новая папка\сtest.docx', 'rb')
        #document = Document(doc)
        #Dictionary = {'radik': "rodion"}
        #fill = Filling.objects.get()
        #for se in fill:
            #data1 = blog_filling.objects.all().order_by('id')
            #ses_id = request.args.get('blog_filling.id')
            #p = document.add_paragraph('radik')
        usdata = False
        usdoc = UserDoc.objects.all()
        for u in usdoc:
            usdata = True
            #doc1 = open(File(u.userFile), 'rb')
            #doc1 = File(u.userFile)
            #print("usdoc", doc1)
            #document = Document(doc1)
            #if(u.user_id == current_user):
            #    doc1 = File(u.userFile)
            #    document = Document(doc1)


        test = UploadFileForm.objects.all()
        for t in test:
            if(usdata == False):
                doc = File(t.file)
                document = Document(doc)
            else:
                doc = File(t.file)
                document = Document(doc)
                old_data = worker.objects.all()

                #for data in old_data:
                    #if(current_user == data.author_id_id):
            #filling_data(request)
            #worker_data(request)
            for p in document.paragraphs:


                #print paragraph.text
                #p.text = inputWname
                inline = p.runs

                for i in range(len(inline)):
                        if Dictionary['inn'] in inline[i].text:
                            text = inline[i].text.replace(Dictionary['inn'], inputInn)
                            inline[i].text = text
                        if Dictionary['ogrn'] in inline[i].text:
                            text = inline[i].text.replace(Dictionary['ogrn'], inputOgrn)
                            inline[i].text = text

                        if Dictionary['orgcity'] in inline[i].text:
                            text = inline[i].text.replace(Dictionary['orgcity'], inputCity)
                            inline[i].text = text
                        if Dictionary['address'] in inline[i].text:
                            text = inline[i].text.replace(Dictionary['address'], inputAddress)
                            inline[i].text = text
                        if Dictionary['orgname'] in inline[i].text:
                            text = inline[i].text.replace(Dictionary['orgname'], inputName)
                            inline[i].text = text

            document.save('C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\' + str(request.user.id) + t.title + '_сtest2.docx')
        directory = "C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\"
        files = os.listdir(directory)
        doc1 = filter(lambda x: x.endswith(str(request.user.id) + t.title + '_сtest2.docx'), files)
        for f in doc1:
            if(finded == True):
                some = UserDoc(id=uid, user=request.user, userFile=f)
                some.save()
            else:
                some = UserDoc.objects.create(user=request.user, userFile=f)
                some.save()
        doc.close()
        return HttpResponseRedirect("/worker")
    #c = {}
    #c.update(csrf(request))
    #return HttpResponseRedirect("/filling")
    return render(request, 'blog/filling_list.html', {"user":cur_user, "message":message})

        #return render(request, 'blog/filling_list.html', {"inn": "123", "ogrn": "2737"})

def organization_data(request):
    #fillingView(request)
    test = UploadFileForm.objects.all()
    #print(usdata)
    global inputInn, inputOgrn, inputCity, inputName, inputAddress, doctype
    #fill = Filling.objects.all()
    cur_user = request.user.id
    for t in test:
        #doc = File(t.file)
        doc = open('C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\' + str(request.user.id) + t.title + '.docx', 'rb')
        document = Document(doc)
        """for s in fill:
            print("name",s.name)
            if(cur_user == s.author_id_id):
                inputInn = s.inn
                inputOgrn = s.ogrn
                inputCity = s.city
                inputAddress = s.address
                inputName = s.name
                print("inn",inputInn)"""
        for p in document.paragraphs:
                #if Dictionary['orgname'] in p.text:
        #print paragraph.text
        #p.text = inputWname
            inline = p.runs
            if p.text == "{{Руководитель}}":
                filling = Filling.objects.filter(author_id_id=request.user.id)
                for f in filling:
                    #print(members_comission)
                    p.insert_paragraph_before("           " + f.chief_fullposition + "                                " + f.chief_name[:1] + "." + f.chief_secondname[:1] + "." + f.chief_surname)
                p.clear()
            for i in range(len(inline)):
                    if Dictionary['inn'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['inn'], str(inputInn))
                        inline[i].text = text
                    if Dictionary['ogrn'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['ogrn'], str(inputOgrn))
                        inline[i].text = text

                    if Dictionary['orgcity'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['orgcity'], inputCity)
                        inline[i].text = text
                    if Dictionary['address'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['address'], inputAddress)
                        inline[i].text = text
                    if Dictionary['orgname'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['orgname'], inputName)
                        inline[i].text = text
        document.save('C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\' + str(request.user.id) + t.title + '.docx')

array_work = dict()
def worker_data(request):
    test = UploadFileForm.objects.all()
    work_data = worker.objects.all()
    cur_user = request.user.id
    global inputWsurname, inputWname, inputLastname, inputPosition, members_comission, array_work
    #print(inputWname)
    for t in test:
        doc = File(t.file)
        #doc = open('C:\\myproject\\document\\' + str(request.user.id) + t.title + '_сtest2.docx')
        document = Document(doc)
        """for w in work_data:
            if(cur_user == w.author_id_id):"""
        for p in document.paragraphs:
            """worker_data = worker.objects.filter(author_id_id=request.user.id)
            for w in worker_data:
                inputWsurname = w.w_surname
                inputWname = w.w_name
                inputLastname = w.w_lastname
                inputPosition = w.position"""
        #print paragraph.text
        #p.text = inputWname
            if p.text == "{{Подписи}}":
                workers = worker.objects.filter(author_id_id=request.user.id)
                for w in workers:
                    #print(members_comission)
                    for named in dict_commision['name']:
                        if named == w.w_name:

                            #p.insert_paragraph_before(w.w_name)
                            p.insert_paragraph_before("             " + w.w_surname + " " + w.w_name + " " + w.w_lastname + "                       " + w.position)


                #p.insert_paragraph_before("Иванов Сидоро Петрович                   директор")
                #p.insert_paragraph_before("Михайлов Миахил Петрович директор")
                #p.insert_paragraph_before("Иванов Петр Петрович директор")
                p.clear()
                #del(document.paragraphs[document.paragraphs.index(p)])
            inline = p.runs

            for i in range(len(inline)):
                    if Dictionary['surname'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['surname'], inputWsurname)
                        inline[i].text = text
                    if Dictionary['name'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['name'], inputWname)
                        inline[i].text = text
                    if Dictionary['second_name'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['second_name'], inputLastname)
                        inline[i].text = text
                    if Dictionary['position'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['position'], inputPosition)
                        inline[i].text = text
            #document.save('C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\' + str(request.user.id) + t.title + '_сtest2.docx')
            document.save('C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\' + str(request.user.id) + t.title + '.docx')


@login_required
def fillingView_change(request):
    current_user = request.user.id
    filling_data = Filling.objects.all()
    #for f in filling_data:
    #if request.method =='POST':
    #    c = {}
    #c.update(csrf(request))
    #return render(request, 'blog/filling_list_change.html', {})
    return HttpResponseRedirect("/filling_change")
        #return render(request, 'blog/filling_list.html', {"inn": "123", "ogrn": "2737"})

#orgid = []

users_worker = {}

@login_required
def org_data(request):
    global inputWsurname, inputWname, inputLastname, inputPosition, doctype, doctitle
    global inputInn, inputOgrn, inputCity, inputAddress, inputName, members_comission
    #count = 0
    if request.POST:
        a = request.POST.get('acting')
        workers = worker.objects.filter(author_id_id=request.user.id)
        for w in workers:
            members_comission = request.POST.get('checkcomm'+str(w.id))
            if members_comission is not None:
                #b = array_work.append(members_comission)
                dict_commision['name'].append(members_comission)
        for w in workers:

            if a == w.w_name:
                inputWsurname = w.w_surname
                inputWname = w.w_name
                inputLastname = w.w_lastname
                inputPosition = w.position
                worker_data(request)
                dict_commision['name'].clear()
            #print(w.id)

            delete = request.POST.get('delete')
            if delete:
                workers = worker.objects.filter(w_name=members_comission).delete()
                return HttpResponseRedirect("/orgdata")
            #for i in workers:
            #for w in workers:
                #users_worker = {w.w_name:1}

                    #print(users_worker)
        fill = Filling.objects.filter(author_id_id=request.user.id)
        for f in fill:
            inputInn = f.inn
            inputOgrn = f.ogrn
            inputCity = f.city
            inputAddress = f.address
            inputName = f.name
            organization_data(request)

        doctemplate = UploadFileForm.objects.all()
        zipp=zipfile.ZipFile("C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\document" + str(request.user.id) + ".zip", mode='w')
        for d in doctemplate:
            #print("yes")
            doctype = request.POST.get('doc' + str(d.id))
            if doctype is not None:
                zipp.write('C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\' + str(request.user.id) + doctype + '.docx')
            #print(doctype)
        #    if doctype == d.title:


        return HttpResponseRedirect("/download")


        #return render(request, 'blog/orgdata.html', {"organization": Filling.objects.all(), "items": worker.objects.all(), "uid":request.user.id, "docx":UploadFileForm.objects.all()})
    current_user = request.user.id
    #zipp.close()
    return render(request, 'blog/orgdata.html', {"organization": Filling.objects.all(), "items": worker.objects.all(), "uid":request.user.id, "docx":UploadFileForm.objects.all()})
    #return render(request, 'blog/orgdata.html', {})

def RTFobject(request):#Функция открытия docx файла
    #document = opendocx('C:\\myproject\\Sun.docx')
    doc = open('C:\myproject\Новая папка\Sun.docx', 'rb')
    document = Document(doc)
    Dictionary = {'inn': "11111"}
    fill = Filling.objects.all()
    #data1 = blog_filling.objects.all().order_by('id')
    #ses_id = request.args.get('blog_filling.id')
    #p = document.add_paragraph('radik')


    document.save('C:\myproject\Новая папка\Sun2.docx')
    return render_to_response('blog/post_list.html')

@login_required
def Worker_change(request):
    if request.method == 'POST':
        inputWname = request.POST.get('w_name')
        inputWsurname = request.POST.get('w_surname')
        inputLastname = request.POST.get('w_lastname')
        inputPosition = request.POST.get('position')
        current_user = request.user.id
        b = worker.objects.create(author_id=request.user, w_name=inputWname, w_surname=inputWsurname, w_lastname=inputLastname, position=inputPosition)
        b.save()
        c = {}
        #c.update(csrf(request))
        #return HttpResponseRedirect("/")

    return render_to_response('blog/worker_change.html')

@login_required
def Worker(request, worker_id=None):
    cur_user = worker() if worker_id is None else worker.objects.get(pk=worker_id)

    # found = False
    # current_user = request.user.id
    # count = False
    # data = worker.objects.all()
    message = ""
    # for f in data:
    #     if(f.author_id_id == current_user):
    #         cur_user = worker.objects.get(pk=f.id)
    #         count = True
    #         found = True
    #         test = f.id
    #
    # if(count == False):
    #     found = True
    #     print("phh")
    #     #return HttpResponseRedirect("/worker_change")
    #     return render_to_response('blog/worker.html')
    #
    # userdoc = UserDoc.objects.all()
    # finded = False
    # for u in userdoc:
    #     if(u.user_id == current_user):
    #         #cur_user = UserDoc.objects.get(pk=u.id)
    #         finded = True
    #         uid = u.id

    if request.method =='POST':
        cur_user.author_id = request.user
        cur_user.w_name = request.POST.get('w_name')
        cur_user.w_surname = request.POST.get('w_surname')
        cur_user.w_lastname = request.POST.get('w_lastname')
        cur_user.position = request.POST.get('position')
        # cur_user[''] = request.user.id
        cur_user.save()
        message = "Данные успешно сохранены"
        #query = worker.objects.get(author_id=author_id)
        # if found:
        #     b = worker(id = test, author_id=request.user, w_name=inputWname, w_surname=inputWsurname, w_lastname=inputLastname, position=inputPosition)
        #     b.save()
        # else:
        #     b = worker.objects.create(author_id=request.user, w_name=inputWname, w_surname=inputWsurname, w_lastname=inputLastname, position=inputPosition)
        #     b.save()
        """directory = "C:\\myproject\\document"
        files = os.listdir(directory)
        doc1 = filter(lambda x: x.endswith('.docx'), files)
        for f in doc1:
            print(f)
            #s = open(f, 'r')
            #s = open(f, 'rb')
            #print(s)

            #s.close()"""




        # usdata = False
        # usdoc = UserDoc.objects.all()
        # for u in usdoc:
        #     if(u.user_id == current_user):
        #         usdata = True
        #
        #
        # test = UploadFileForm.objects.all()
        # for t in test:
        #     doc = File(t.file)
        #     document = Document(doc)
        #     if(usdata == False):
        #         doc = File(t.file)
        #         document = Document(doc)
        #     else:
        #         #print("tis")
        #         #doc = open('C:\\myproject\\document\\'+ str(request.user.id) + t.title + '_сtest2.docx', 'rb')
        #         #document = Document(doc)
        #         doc = File(t.file)
        #         document = Document(doc)
        #     for p in document.paragraphs:
        #         if Dictionary['position'] in p.text:
        #             #print paragraph.text
        #             #p.text = inputWname
        #             inline = p.runs
        #
        #             for i in range(len(inline)):
        #                     if Dictionary['surname'] in inline[i].text:
        #                         text = inline[i].text.replace(Dictionary['surname'], inputWsurname)
        #                         inline[i].text = text
        #                     if Dictionary['name'] in inline[i].text:
        #                         text = inline[i].text.replace(Dictionary['name'], inputWname)
        #                         inline[i].text = text
        #
        #                     if Dictionary['second_name'] in inline[i].text:
        #                         text = inline[i].text.replace(Dictionary['second_name'], inputLastname)
        #                         inline[i].text = text
        #                     if Dictionary['position'] in inline[i].text:
        #                         text = inline[i].text.replace(Dictionary['position'], inputPosition)
        #                         inline[i].text = text
        #     document.save('C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\' + str(request.user.id) + t.title + '_сtest2.docx')
        #     """if(finded == True):
        #         some = UserDoc(id=uid, user=request.user, userFile=doc)
        #         some.save()
        #     else:
        #         some = UserDoc.objects.create(user=request.user, userFile=doc)
        #         some.save()"""
        #     directory = 'C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\'
        #     files = os.listdir(directory)
        #     print(files)
        #     doc1 = filter(lambda x: x.endswith(str(request.user.id) + t.title + '_сtest2.docx'), files)
        #     for f in doc1:
        #         print(f)
        #         if(finded == True):
        #             some = UserDoc(id=uid, user=request.user, userFile=f)
        #             some.save()
        #         else:
        #             some = UserDoc.objects.create(user=request.user, userFile=f)
        #             some.save()
        #     #doc.close()
        #         #print("userFile", some.userFile)
        #     #some = UserDoc(user=request.user, userFile=)
        #     #print(some.userFile)
        #
        #     #b = UserDoc.objects.create(user=request.user, userFile=text)
        #     #some.save()
        #
        #     #newuserdoc = UserDoc.userFile(userFile = document)
        #     #newuserdoc.save()
        #     #b = UserDoc.objects.create(user=request.user, userFile=sav)
        # fill = Filling.objects.all()
        # #data1 = blog_filling.objects.all().order_by('id')1
        # #ses_id = request.args.get('blog_filling.id')
        # #p = document.add_paragraph('radik')
        #
        # """for p in document.paragraphs:
        #     if Dictionary['position'] in p.text:
        #         #print paragraph.text
        #         #p.text = inputWname
        #         inline = p.runs
        #
        #         for i in range(len(inline)):
        #                 if Dictionary['surname'] in inline[i].text:
        #                     text = inline[i].text.replace(Dictionary['surname'], inputWsurname)
        #                     inline[i].text = text
        #                 if Dictionary['name'] in inline[i].text:
        #                     text = inline[i].text.replace(Dictionary['name'], inputWname)
        #                     inline[i].text = text
        #
        #                 if Dictionary['second_name'] in inline[i].text:
        #                     text = inline[i].text.replace(Dictionary['second_name'], inputLastname)
        #                     inline[i].text = text
        #                 if Dictionary['position'] in inline[i].text:
        #                     text = inline[i].text.replace(Dictionary['position'], inputPosition)
        #                     inline[i].text = text
        #
        # document.save('C:\myproject\Новая папка\сtest2.docx')"""
    #return HttpResponseRedirect("/worker")
    return render(request, 'blog/worker.html', {"user": cur_user, "message": message})

def upload_file(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            instance = ModelWithFileField(file_field=request.FILES['file'])
            instance.save()
            return HttpResponseRedirect('/')
    else:
        form = UploadFileForm()
        print(form.title)

    return render(request, 'blog/upload.html',{"form":form})
