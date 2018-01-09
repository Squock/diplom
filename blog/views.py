from django.forms import ModelForm, TextInput
from django.shortcuts import render
from django.contrib import auth
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import UserCreationForm
from django.shortcuts import render_to_response
from django.urls import reverse, reverse_lazy
from django.views.generic import UpdateView
from docx import Document
from django.contrib.auth.models import User
from blog.models import Organization, Userdata, Worker, UploadFileForm
from django.http import HttpResponseRedirect, HttpResponse
import os
from django.views.decorators.csrf import requires_csrf_token
from django.core.files import File
import zipfile
from collections import defaultdict

Dictionary = {'name': 'Имя',
              'surname': 'Фамилия',
              'second_name': 'Отчество',
              'position': 'Должность',
              'inn': 'inn',
              'ogrn': 'ogrn',
              'orgname': 'orgname',
              'address': 'address',
              'orgcity': 'city'}
dict_commision = defaultdict(list)
# dict_commision['name']
idd = []
users_worker = {}


def logout(request):
    auth.logout(request)
    # Перенаправление на страницу.
    return HttpResponseRedirect("/")


@login_required
def contact(request):
    return render(request, 'blog/contact.html', {})


@login_required
def cabinet(request):
    message = " "

    current_user = request.user.id
    count = False
    data = Userdata.objects.all()
    for f in data:
        if f.user_id == current_user:
            cur_user = Userdata.objects.get(pk=f.id)
            count = True
    if count is False:
        print("phh")
        return HttpResponseRedirect("/userdata")

    if request.method == "POST":
        cur_user.name = request.POST["name"]
        cur_user.surname = request.POST["surname"]
        cur_user.email = request.POST["email"]
        cur_user.number = request.POST["number"]
        cur_user.second_name = request.POST["second_name"]
        cur_user.save()
        message = "Данные успешно сохранены"
    return render(request, 'blog/cabinet.html', {"user": cur_user, "message": message})


@login_required
def pass_set(request):
    message = ""
    if request.method == "POST":
        u = User.objects.get(username=request.user)
        password = request.POST["password2"]
        password2 = request.POST["password"]
        if password == password2:
            u.set_password(password)
            u.save()
            print(u)
        else:
            message = "Данные не совпадают"
    return render(request, 'blog/pass_set.html', {"message": message})


@login_required
def post_list(request):
    filename = UploadFileForm.objects.all()
    for f in filename:
        docfile = f.file

    if request.method == 'POST':
        return HttpResponse(docfile, mimetype='application/octet-stream')

    return render(request, 'blog/post_list.html', {})


def download(request):
    zipfile = open("C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\document" + str(
        request.user.id) + ".zip", 'rb')
    response = HttpResponse(content=zipfile)
    response['Content-Type'] = 'application/zip'
    response['Content-Disposition'] = "attachment; filename=document.zip"
    return response


"""def login(request):
    username = request.POST['username']
    password = request.POST['password']
    user = auth.authenticate(username=username, password=password)
    if user is not None and user.is_active:
        # Правильный пароль и пользователь "активен"
        auth.login(request, user)
        # Перенаправление на "правильную" страницу
        return HttpResponseRedirect("/account/loggedin/")
    else:
        # Отображение страницы с ошибкой
        return HttpResponseRedirect("/account/invalid/")"""

def registrationView(request):  # Функция регистрации при помощи использование форм Django
    args = {}
    args['form'] = UserCreationForm()
    if request.method == 'POST':
        newuser_form = UserCreationForm(request.POST)
        if newuser_form.is_valid():
            newuser_form.save()
            newuser = auth.authenticate(username=newuser_form.cleaned_data['username'],
                                        password=newuser_form.cleaned_data['password2'])
        else:
            args['form'] = newuser_form
    return render_to_response('registration/registration.html', args)


@requires_csrf_token
def userview(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        surname = request.POST.get('surname')
        second_name = request.POST.get('second_name')
        email = request.POST.get('email')
        number = request.POST.get('number')
        b = Userdata.objects.create(user=request.user, name=name, surname=surname,
                                    second_name=second_name, email=email, number=number)
        b.save()
        return HttpResponseRedirect("/")
    return render(request, 'registration/user.html', {})


@login_required
def fillingView(request):
    message = ""
    find = False
    cur_org = Organization.objects.filter(author=request.user)
    print("s", cur_org)
    if Organization.objects.filter(author=request.user).exists():
        cur_org = Organization.objects.get(author=request.user)
        find = True

    if request.method == 'POST':
        inn = request.POST.get('inn')
        ogrn = request.POST.get('ogrn')
        city = request.POST.get('city')
        inputAddress = request.POST.get('fact_address')
        org_name = request.POST.get('name')
        message = "Данные успешно сохранены"

        if find is True:
            b = Organization(id=cur_org.id, author=request.user, inn=inn, ogrn=ogrn, name=org_name, city=city)
            b.save()
        else:
            b = Organization.objects.create(author=request.user, inn=inn, ogrn=ogrn, name=org_name, city=city)
            b.save()

    return render(request, 'blog/filling_list.html', {"user": cur_org, "message": message})


class OrganizationForm(ModelForm):
    class Meta:
        model = Organization
        fields = ['inn', 'ogrn', 'name', 'city', 'fact_address', 'reg_address', 'post_address', 'chief']
        widgets = {
            'inn': TextInput(attrs={'class': 'form-control'}),
            'ogrn': TextInput(attrs={'class': 'form-control'}),
            'name': TextInput(attrs={'class': 'form-control'}),
            'city': TextInput(attrs={'class': 'form-control'}),
            'fact_address': TextInput(attrs={'class': 'form-control'}),
            'reg_address': TextInput(attrs={'class': 'form-control'}),
            'post_address': TextInput(attrs={'class': 'form-control'}),
            #'chief': TextInput(attrs={'class': 'form-control'}),
        }


class OrganizationUpdate(UpdateView):
    form_class = OrganizationForm
    model = Organization
    template_name = 'blog/organization.html'
    success_url = reverse_lazy('index')


class WorkerForm(ModelForm):
    class Meta:
        model = Worker
        fields = ['name', 'surname', 'second_name', 'position']
        widgets = {
            'name': TextInput(attrs={'class': 'form-control'}),
            'surname': TextInput(attrs={'class': 'form-control'}),
            'second_name': TextInput(attrs={'class': 'form-control'}),
            'position': TextInput(attrs={'class': 'form-control'}),
        }


class WorkerUpdate(UpdateView):
    form_class = WorkerForm
    model = Worker
    template_name = 'blog/worker_update.html'
    success_url = reverse_lazy('index')


def organization_data(request):
    test = UploadFileForm.objects.all()
    for t in test:
        doc = open('C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\' + str(
            request.user.id) + t.title + '.docx', 'rb')
        document = Document(doc)
        for p in document.paragraphs:
            inline = p.runs
            organization = Organization.objects.filter(author=request.user)
            for org in organization:
                for i in range(len(inline)):
                    if Dictionary['inn'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['inn'], org.inn)
                        inline[i].text = text
                    if Dictionary['ogrn'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['ogrn'], org.ogrn)
                        inline[i].text = text

                    if Dictionary['orgcity'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['orgcity'], org.city)
                        inline[i].text = text
                    # if Dictionary['address'] in inline[i].text:
                    #    text = inline[i].text.replace(Dictionary['address'], inputAddress)
                    #    inline[i].text = text
                    if Dictionary['orgname'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['orgname'], org.name)
                        inline[i].text = text
        document.save('C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\' + str(
            request.user.id) + t.title + '.docx')


array_work = dict()


def worker_data(request):
    # global members_comission
    test = UploadFileForm.objects.all()
    org = Organization.objects.get(author=request.user)
    for t in test:
        doc = File(t.file)
        document = Document(doc)
        for p in document.paragraphs:
            if p.text == "{{Подписи}}":
                workers = Worker.objects.filter(organization_id=org.id)
                for w in workers:
                    for named in dict_commision['name']:
                        if named == w.name:
                            p.insert_paragraph_before(
                                "             " + w.surname + " " + w.name + " " + w.second_name + " " + w.position)
                p.clear()
            inline = p.runs
            workers = Worker.objects.filter(organization_id=org.id)
            for w in workers:
                for i in range(len(inline)):
                    if Dictionary['surname'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['surname'], w.surname)
                        inline[i].text = text
                    if Dictionary['name'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['name'], w.name)
                        inline[i].text = text
                    if Dictionary['second_name'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['second_name'], w.second_name)
                        inline[i].text = text
                    if Dictionary['position'] in inline[i].text:
                        text = inline[i].text.replace(Dictionary['position'], w.position)
                        inline[i].text = text
                document.save('C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\' + str(
                    request.user.id) + t.title + '.docx')


@login_required
def org_data(request):
    user_orgs = [org for org in Organization.objects.filter(author=request.user)]
    if len(user_orgs) == 0:
        cur_org = Organization()
    else:
        cur_org = user_orgs[0]

    if request.method == 'POST':
        cur_org.inn = request.POST.get('inn')
        cur_org.ogrn = request.POST.get('ogrn')
        cur_org.city = request.POST.get('city')
        cur_org.address = request.POST.get('address')
        cur_org.name = request.POST.get('name')
        cur_org.save()
    return render(request, 'blog/orgdata.html',
                  {"user": cur_org})


@login_required
def choose_worker(request):
    user_orgs = [org for org in Organization.objects.filter(author=request.user)]
    if len(user_orgs) == 0:
        cur_org = Organization()
    else:
        cur_org = user_orgs[0]

    organization = Organization.objects.filter(author=request.user)
    for org in organization:
        org_id = org.id
        workers = Worker.objects.filter(organization_id=org_id)
        acting = request.POST.get('acting')
        for w in workers:
            members_comission = request.POST.get('checkcomm' + str(w.id))
            if members_comission is not None:
                dict_commision['name'].append(members_comission)
                work = Worker.objects.filter(id=acting)
                for w in work:
                   # worker_data(request)
                   # organization_data(request)
                    dict_commision['name'].clear()
                    return HttpResponseRedirect("/docx")
   # if request.POST:


        #if members_comission is not None:
           # dict_commision['name'].append(members_comission)

    return render(request, 'blog/worker_data.html',
                  {"items": Worker.objects.all(),
                   "oid": cur_org,
                   "docx": UploadFileForm.objects.all()})


@login_required
def download_docx(request):
    if request.POST:
        doctemplate = UploadFileForm.objects.all()
        zipp = zipfile.ZipFile(
            "C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\document" + str(
                request.user.id) + ".zip",
            mode='w')
        for d in doctemplate:
            doctype = request.POST.get('doc' + str(d.id))
            print(doctype)
            if doctype is not None:
                    zipp.write('C:\\myproject\\Новая папка\\django 1.8\\mysite\\media\\uploads\\' + str(
                        request.user.id) + d.title + '.docx')
        return HttpResponseRedirect("/download")
    return render(request, 'blog/download_docx.html', {"docx": UploadFileForm.objects.all()})

@login_required
def worker(request, organization_id=None):
    cur_user = Worker() if organization_id is None else Worker.objects.get(pk=organization_id)
    message = ""
    if request.method == 'POST':
        cur_user.name = request.POST.get('name')
        cur_user.surname = request.POST.get('surname')
        cur_user.second_name = request.POST.get('second_name')
        cur_user.position = request.POST.get('position')
        cur_user.organization_id = request.Worker.organization_id
        cur_user.save()
        message = "Данные успешно сохранены"

    return render(request, 'blog/worker.html', {"user": cur_user, "message": message})
